#!/usr/bin/env python3
"""
Mayank Goyal -- AI Consulting Resume Generator

Produces an ATS-safe Word resume that follows the spec in
Mayank_Resume_Generator_Prompt.md (Sections 6, 9, 10, 12).

Design choices:
- python-docx for full programmatic control of fonts, margins, and spacing.
- No tables, text boxes, columns, headers/footers, images, or page numbers.
- Plain paragraphs + hanging-indent bullets so ATS parses everything as text.
- Hyperlinks inserted via the standard OOXML w:hyperlink workaround.
- Optional LibreOffice headless step converts the .docx to .pdf without
  re-authoring the layout.

Run:
    python3 generate_resume.py                       # Master, today's date
    python3 generate_resume.py --variant Tiger       # name the variant
    python3 generate_resume.py --skip-pdf            # only the .docx
"""

from __future__ import annotations

import argparse
import datetime as dt
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)
ACCENT_HEX = "1F4E79"
BODY_FONT = "Calibri"
BODY_PT = Pt(10)
HEADING_PT = Pt(11)
NAME_PT = Pt(16)
TITLE_PT = Pt(11)
LINE_SPACING = 1.15

SCRIPT_DIR = Path(__file__).resolve().parent
HYPERLINK_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


# ---------- low-level helpers ----------

def _spacing(p, before_pt: float = 0, after_pt: float = 0,
             line_spacing: float = LINE_SPACING) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = line_spacing


def _run(paragraph, text: str, *, bold: bool = False, italic: bool = False,
         size: Pt = BODY_PT, color: RGBColor | None = None):
    r = paragraph.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = size
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def _hyperlink(paragraph, url: str, text: str, *, size_pt: int = 10) -> None:
    """Append a clickable hyperlink to the paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, HYPERLINK_REL, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    rPr.append(fonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size_pt * 2))  # half-points
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT_HEX)
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _set_margins(doc: Document, inches: float = 0.6) -> None:
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _configure_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_PT
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    _spacing(p, before_pt=8, after_pt=2)
    _run(p, text.upper(), bold=True, size=HEADING_PT, color=ACCENT_COLOR)

    # Thin accent rule under each heading -- a paragraph border, not a table.
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _bullet(doc: Document, segments):
    """Bullet paragraph with hanging indent and unicode marker.

    `segments` may be a plain string or a list whose entries are either
    plain strings or (text, bold) tuples.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.22)
    pf.first_line_indent = Inches(-0.22)
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = LINE_SPACING

    _run(p, "•  ")  # bullet glyph + two spaces

    if isinstance(segments, str):
        _run(p, segments)
    else:
        for seg in segments:
            if isinstance(seg, str):
                _run(p, seg)
            else:
                text, bold = seg
                _run(p, text, bold=bold)
    return p


# ---------- section builders ----------

def add_header(doc: Document) -> None:
    p = doc.add_paragraph()
    _spacing(p, after_pt=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "MAYANK GOYAL", bold=True, size=NAME_PT, color=ACCENT_COLOR)

    p = doc.add_paragraph()
    _spacing(p, after_pt=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Senior Data & AI Solutions Consultant", size=TITLE_PT)

    p = doc.add_paragraph()
    _spacing(p, after_pt=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "+91 96466-84712  |  ")
    _hyperlink(p, "mailto:goyalmayank48@gmail.com", "goyalmayank48@gmail.com")
    _run(p, "  |  ")
    _hyperlink(p, "https://linkedin.com/in/digimayank", "linkedin.com/in/digimayank")
    _run(p, "  |  ")
    _hyperlink(p, "https://mayankgoyal.ai", "mayankgoyal.ai")

    p = doc.add_paragraph()
    _spacing(p, after_pt=4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(
        p,
        "Gurugram, India  |  Open to Remote India & Relocation "
        "(Bengaluru / Mumbai / Hyderabad / Pune / Dubai)",
        italic=True,
    )


def add_summary(doc: Document) -> None:
    _section_heading(doc, "Professional Summary")
    p = doc.add_paragraph()
    _spacing(p, after_pt=4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(
        p,
        "Senior Consultant with 5+ years architecting enterprise BI for a "
        "Fortune 100 FMCG portfolio across 60+ markets on a 300GB+ Azure data "
        "lake. Deep on Power BI, DAX, SSAS, and Azure Analysis Services — now "
        "shipping production Generative AI on Python, Anthropic Claude, "
        "LangChain, and Azure OpenAI for senior client engagements.",
    )


def add_core_skills(doc: Document) -> None:
    _section_heading(doc, "Core Skills")
    rows = [
        (
            "Data & Analytics",
            "Power BI · DAX · SQL · Teradata · MS SSAS · "
            "Azure Analysis Services · Tabular Editor · SSMS · "
            "Tabular Models · Semantic Layer · Data Modeling · "
            "Performance Tuning · ETL/ELT · Data Pipelines",
        ),
        (
            "GenAI & ML",
            "Python · Anthropic Claude · OpenAI GPT · Large Language "
            "Models (LLMs) · LangChain · Retrieval-Augmented Generation "
            "(RAG) · Vector Databases (Pinecone, ChromaDB, FAISS) · "
            "Embeddings · Semantic Search · Prompt Engineering · "
            "AI Agents · Tool Use · LLM Evaluation · Streamlit · "
            "FastAPI · Hugging Face · Responsible AI",
        ),
        (
            "Cloud & DevOps",
            "Microsoft Azure · Azure OpenAI · Azure AI Studio · "
            "Azure Data Lake · Azure DevOps · MSAL (Service Principal "
            "Auth) · Power BI REST API · Git · GitHub · "
            "CI/CD · MLOps · Docker · Microservices · "
            "API Development · Microsoft Power Automate",
        ),
        (
            "Consulting & Leadership",
            "Solution Architecture · Stakeholder Management · Client "
            "Engagement · Cross-functional Team Leadership · Agile / "
            "Scrum Delivery · Technical Mentoring · POC to Production "
            "· Business Storytelling · AI Governance",
        ),
    ]
    for label, body in rows:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing = LINE_SPACING
        _run(p, f"{label}:  ", bold=True)
        _run(p, body)


def add_featured_project(doc: Document) -> None:
    _section_heading(doc, "Featured Project")

    p = doc.add_paragraph()
    _spacing(p, after_pt=1)
    _run(p, "PowerBI Insight Co-pilot", bold=True)
    _run(p, "  —  Open-Source GenAI Project  ")
    _run(p, "(Jun 2026)", italic=True)

    p = doc.add_paragraph()
    _spacing(p, after_pt=2)
    _run(p, "Live: ", bold=True)
    _hyperlink(
        p,
        "https://mayankgoyal.ai/projects/powerbi-copilot",
        "mayankgoyal.ai/projects/powerbi-copilot",
    )
    _run(p, "    Code: ", bold=True)
    _hyperlink(
        p,
        "https://github.com/digimayank/powerbi-insight-copilot",
        "github.com/digimayank/powerbi-insight-copilot",
    )

    p = doc.add_paragraph()
    _spacing(p, after_pt=3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(
        p,
        "Production-grade natural language interface for Power BI semantic "
        "models. Users query in English; Anthropic Claude generates DAX, "
        "executes against the Power BI REST API, and returns narrative "
        "insights with source attribution — reducing analyst turnaround "
        "time on routine business queries by an estimated 60%.",
    )

    _bullet(doc, [
        ("Stack: ", True),
        "Python · Anthropic Claude SDK · LangChain patterns · "
        "Power BI REST API · Azure AD (MSAL service principal auth) · "
        "Streamlit · Streamlit Cloud",
    ])
    _bullet(doc, [
        ("Architecture: ", True),
        "3-stage orchestrator (natural language → DAX generation → "
        "execution → narration) with LLM-driven DAX self-correction, "
        "schema-aware prompting, conversation memory for follow-up queries, "
        "and production-ready error handling and retry logic",
    ])
    _bullet(doc, [
        ("Differentiator: ", True),
        "Engineered for enterprise constraints — single semantic layer, "
        "service principal authentication, schema-bounded generation, and "
        "end-to-end traceability suitable for governed client environments",
    ])


def add_experience(doc: Document) -> None:
    _section_heading(doc, "Professional Experience")

    # --- Accenture ---
    p = doc.add_paragraph()
    _spacing(p, after_pt=0)
    _run(p, "Accenture Solutions Pvt Ltd", bold=True)
    _run(p, "  —  Senior Analyst, Data & Analytics")

    p = doc.add_paragraph()
    _spacing(p, after_pt=3)
    _run(p, "Gurugram, India (Hybrid)   |   Jul 2024 – Present", italic=True)

    _bullet(doc,
        "Architect BI infrastructure for a Fortune 100 FMCG client's "
        "multi-billion-dollar global product portfolio across 60+ markets — "
        "own the semantic layer on a 300GB+ Azure-hosted data lake; partner with "
        "senior client stakeholders to convert strategic decisions into "
        "measurable topline outcomes")
    _bullet(doc,
        "Engineer Azure Analysis Services tabular models on one of the largest "
        "production BI data lakes in the FMCG industry; implemented incremental "
        "refresh and partition strategies enabling sub-second query response "
        "at enterprise scale")
    _bullet(doc,
        "Coordinate a 40+ member cross-functional delivery team spanning data "
        "engineering, BI development, QA, and client stakeholders across 3 "
        "continents; established a daily delivery rhythm that reduced critical "
        "bug resolution time by 50%")
    _bullet(doc,
        "Lead technical mentorship for new joiners on advanced DAX, performance "
        "tuning, and semantic model architecture; built reusable internal "
        "frameworks adopted across multiple client engagements")

    # --- PwC ---
    p = doc.add_paragraph()
    _spacing(p, before_pt=4, after_pt=0)
    _run(p, "PwC India", bold=True)
    _run(p, "  —  Associate (Power BI Consultant)")

    p = doc.add_paragraph()
    _spacing(p, after_pt=3)
    _run(p, "Gurugram, India (Remote)   |   Mar 2021 – Jul 2024", italic=True)

    _bullet(doc,
        "Senior analyst on the Global Net Revenue Management (NRM) initiative "
        "for a leading Fortune 100 FMCG client — owned BI reporting "
        "infrastructure informing pricing, promotional, and product-mix "
        "decisions across the global portfolio")
    _bullet(doc,
        "Architected migration of 40+ business-critical dashboards from Tableau "
        "to Power BI; redesigned the semantic layer and DAX measure architecture, "
        "reducing report load latency by 75% on 300GB+ tabular models")
    _bullet(doc,
        "Translated complex C-suite business requirements into solution "
        "architectures spanning data modeling, integration patterns, and "
        "visualization workflows; partnered with clients on roadmaps, "
        "prioritization, and delivery sequencing")
    _bullet(doc,
        "Conducted 50+ technical interviews evaluating Power BI candidates and "
        "led onboarding of new team members through structured technical "
        "enablement programs")


def add_education(doc: Document) -> None:
    _section_heading(doc, "Education")
    p = doc.add_paragraph()
    _spacing(p, after_pt=0)
    _run(p, "Bachelor of Technology — Computer Science Engineering", bold=True)
    p = doc.add_paragraph()
    _spacing(p, after_pt=2)
    _run(
        p,
        "Giani Zail Singh College of Engineering & Technology, Punjab Technical "
        "University, Bathinda   |   2016 – 2020",
        italic=True,
    )


def add_certifications(doc: Document) -> None:
    _section_heading(doc, "Certifications")
    _bullet(doc, [
        ("Microsoft Certified: Azure AI Engineer Associate (AI-102)", True),
        "  —  In Progress, target Jul 2026",
    ])
    _bullet(doc, [
        ("Analyzing and Visualizing Data with Microsoft Power BI", True),
        "  —  Microsoft (Apr 2021)",
    ])
    _bullet(doc, [
        ("Marketing Analytics", True),
        "  —  Udacity (May 2020)",
    ])


def add_languages(doc: Document) -> None:
    _section_heading(doc, "Languages")
    p = doc.add_paragraph()
    _spacing(p, after_pt=2)
    _run(
        p,
        "English (professional working proficiency)  ·  Hindi (native)  "
        "·  Punjabi (native)",
    )


# ---------- orchestration ----------

def build_resume(output_path: Path) -> Path:
    doc = Document()
    _set_margins(doc, 0.6)
    _configure_normal_style(doc)

    add_header(doc)
    add_summary(doc)
    add_core_skills(doc)
    add_featured_project(doc)
    add_experience(doc)
    add_education(doc)
    add_certifications(doc)
    add_languages(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def convert_to_pdf(docx_path: Path) -> Path | None:
    if shutil.which("soffice") is None:
        print("WARN: soffice not found on PATH; skipping PDF conversion.",
              file=sys.stderr)
        return None
    out_dir = docx_path.parent
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(out_dir), str(docx_path)],
            check=True,
            capture_output=True,
            timeout=180,
        )
    except subprocess.CalledProcessError as exc:
        print(f"WARN: PDF conversion failed (exit {exc.returncode}): "
              f"{exc.stderr.decode(errors='replace')}", file=sys.stderr)
        return None
    except subprocess.TimeoutExpired:
        print("WARN: PDF conversion timed out.", file=sys.stderr)
        return None
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    return pdf_path if pdf_path.exists() else None


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--variant", default="Master",
                        help="Variant label baked into the filename "
                             "(Master, AI_Consulting_India, Big4_India, Tiger, ...)")
    parser.add_argument("--outdir", default=str(SCRIPT_DIR / "outputs"),
                        help="Output directory (default: ./outputs)")
    parser.add_argument("--date", default=None,
                        help="YYYYMMDD stamp; defaults to today")
    parser.add_argument("--skip-pdf", action="store_true",
                        help="Skip the LibreOffice PDF conversion step")
    args = parser.parse_args()

    date_str = args.date or dt.date.today().strftime("%Y%m%d")
    out_dir = Path(args.outdir)
    docx_path = out_dir / f"Mayank_Goyal_Resume_{args.variant}_{date_str}.docx"

    build_resume(docx_path)
    print(f"Wrote {docx_path}")

    if not args.skip_pdf:
        pdf_path = convert_to_pdf(docx_path)
        if pdf_path is not None:
            print(f"Wrote {pdf_path}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
